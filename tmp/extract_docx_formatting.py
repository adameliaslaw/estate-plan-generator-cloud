"""
Extract detailed formatting information from all .docx sample files.
Outputs paragraph styles, runs (bold/italic/underline/font/size),
alignment, spacing, indentation, tables, headers/footers, page setup, etc.
"""
import os
import json
from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

SAMPLES_DIR = r"c:\estate-plan-generator\samples"

ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
    WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
    WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
    None: "INHERIT/DEFAULT",
}

def emu_to_inches(val):
    if val is None:
        return None
    return round(val / 914400, 3)

def pt_val(val):
    if val is None:
        return None
    return float(val.pt) if hasattr(val, 'pt') else val

def extract_paragraph(p, idx):
    pf = p.paragraph_format
    info = {
        "index": idx,
        "style": p.style.name if p.style else None,
        "alignment": ALIGN_MAP.get(pf.alignment, str(pf.alignment)),
        "space_before_pt": pt_val(pf.space_before),
        "space_after_pt": pt_val(pf.space_after),
        "line_spacing": pt_val(pf.line_spacing) if pf.line_spacing else None,
        "left_indent_in": emu_to_inches(pf.left_indent) if pf.left_indent else None,
        "right_indent_in": emu_to_inches(pf.right_indent) if pf.right_indent else None,
        "first_line_indent_in": emu_to_inches(pf.first_line_indent) if pf.first_line_indent else None,
        "keep_together": pf.keep_together,
        "keep_with_next": pf.keep_with_next,
        "page_break_before": pf.page_break_before,
        "text_preview": p.text[:120] if p.text else "",
        "runs": [],
    }
    for r in p.runs:
        rf = r.font
        run_info = {
            "text": r.text[:80] if r.text else "",
            "bold": rf.bold,
            "italic": rf.italic,
            "underline": rf.underline,
            "font_name": rf.name,
            "font_size_pt": float(rf.size.pt) if rf.size else None,
            "all_caps": rf.all_caps,
            "small_caps": rf.small_caps,
            "color_rgb": str(rf.color.rgb) if rf.color and rf.color.rgb else None,
        }
        info["runs"].append(run_info)
    return info

def extract_table(t, idx):
    rows_data = []
    for ri, row in enumerate(t.rows):
        cells_data = []
        for ci, cell in enumerate(row.cells):
            cells_data.append({
                "col": ci,
                "text_preview": cell.text[:100] if cell.text else "",
                "paragraphs_count": len(cell.paragraphs),
            })
        rows_data.append({"row": ri, "cells": cells_data})
    return {"table_index": idx, "rows": len(t.rows), "cols": len(t.columns), "data": rows_data}

def extract_section(sec, idx):
    return {
        "section_index": idx,
        "page_width_in": emu_to_inches(sec.page_width),
        "page_height_in": emu_to_inches(sec.page_height),
        "left_margin_in": emu_to_inches(sec.left_margin),
        "right_margin_in": emu_to_inches(sec.right_margin),
        "top_margin_in": emu_to_inches(sec.top_margin),
        "bottom_margin_in": emu_to_inches(sec.bottom_margin),
        "header_distance_in": emu_to_inches(sec.header_distance),
        "footer_distance_in": emu_to_inches(sec.footer_distance),
        "different_first_page_header": sec.different_first_page_header_footer,
    }

def extract_doc(filepath):
    doc = Document(filepath)
    result = {
        "filename": os.path.basename(filepath),
        "sections": [],
        "paragraphs": [],
        "tables": [],
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
    }
    
    for i, sec in enumerate(doc.sections):
        result["sections"].append(extract_section(sec, i))

    for i, p in enumerate(doc.paragraphs):
        result["paragraphs"].append(extract_paragraph(p, i))

    for i, t in enumerate(doc.tables):
        result["tables"].append(extract_table(t, i))

    return result

def main():
    output_dir = r"c:\estate-plan-generator\tmp"
    os.makedirs(output_dir, exist_ok=True)
    
    files = sorted([f for f in os.listdir(SAMPLES_DIR) if f.endswith('.docx')])
    print(f"Found {len(files)} docx files\n")
    
    for fname in files:
        fpath = os.path.join(SAMPLES_DIR, fname)
        print(f"Processing: {fname}")
        data = extract_doc(fpath)
        
        out_path = os.path.join(output_dir, fname.replace('.docx', '_formatting.json'))
        with open(out_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, default=str)
        print(f"  -> {out_path}")
        
        # Also output a readable text summary
        txt_path = os.path.join(output_dir, fname.replace('.docx', '_content.txt'))
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(f"=== {fname} ===\n\n")
            for sec in data["sections"]:
                f.write(f"[Section {sec['section_index']}] "
                        f"Page: {sec['page_width_in']}x{sec['page_height_in']}in, "
                        f"Margins: L={sec['left_margin_in']} R={sec['right_margin_in']} "
                        f"T={sec['top_margin_in']} B={sec['bottom_margin_in']}in\n")
            f.write("\n--- PARAGRAPHS ---\n\n")
            for p in data["paragraphs"]:
                style = p["style"] or "None"
                align = p["alignment"] or "?"
                text = p["text_preview"]
                f.write(f"[P{p['index']}] Style={style} | Align={align}")
                if p["space_before_pt"]:
                    f.write(f" | SpBefore={p['space_before_pt']}pt")
                if p["space_after_pt"]:
                    f.write(f" | SpAfter={p['space_after_pt']}pt")
                if p["line_spacing"]:
                    f.write(f" | LineSpacing={p['line_spacing']}")
                if p["left_indent_in"]:
                    f.write(f" | LeftInd={p['left_indent_in']}in")
                if p["first_line_indent_in"]:
                    f.write(f" | FirstLine={p['first_line_indent_in']}in")
                if p["page_break_before"]:
                    f.write(" | PAGE_BREAK")
                f.write(f"\n  Text: {text}\n")
                for r in p["runs"]:
                    attrs = []
                    if r["bold"]: attrs.append("BOLD")
                    if r["italic"]: attrs.append("ITALIC")
                    if r["underline"]: attrs.append("UNDERLINE")
                    if r["all_caps"]: attrs.append("ALL_CAPS")
                    if r["small_caps"]: attrs.append("SMALL_CAPS")
                    if r["font_name"]: attrs.append(f"Font={r['font_name']}")
                    if r["font_size_pt"]: attrs.append(f"Size={r['font_size_pt']}pt")
                    if r["color_rgb"]: attrs.append(f"Color=#{r['color_rgb']}")
                    attr_str = ", ".join(attrs) if attrs else "default"
                    f.write(f"    Run[{attr_str}]: \"{r['text']}\"\n")
                f.write("\n")
            
            if data["tables"]:
                f.write("\n--- TABLES ---\n\n")
                for t in data["tables"]:
                    f.write(f"[Table {t['table_index']}] {t['rows']}x{t['cols']}\n")
                    for row in t["data"]:
                        for cell in row["cells"]:
                            f.write(f"  [{row['row']},{cell['col']}]: {cell['text_preview']}\n")
                    f.write("\n")
        
        print(f"  -> {txt_path}")
    
    print("\nDone!")

if __name__ == "__main__":
    main()
